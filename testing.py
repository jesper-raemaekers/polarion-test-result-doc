from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree
import xml.etree.ElementTree as ET
from io import StringIO, BytesIO

from docx.oxml.text.paragraph import CT_P
from docx.oxml.simpletypes import ST_String
from docx.oxml import *
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OptionalAttribute,
    RequiredAttribute,
    ZeroOrMore,
    ZeroOrOne,
    OneAndOnlyOne,
)

class CT_Tag(BaseOxmlElement):
    """
    Used for ``<w:b>``, ``<w:i>`` elements and others, containing a bool-ish
    string in its ``val`` attribute, xsd:boolean plus 'on' and 'off'.
    """
    val = OptionalAttribute('w:val', ST_String, default=True)

class CT_Alias(BaseOxmlElement):
    """
    Used for ``<w:b>``, ``<w:i>`` elements and others, containing a bool-ish
    string in its ``val`` attribute, xsd:boolean plus 'on' and 'off'.
    """
    val = OptionalAttribute('w:val', ST_String, default=True)

class CT_BlockProperties(BaseOxmlElement):
    """
    Used for ``<w:b>``, ``<w:i>`` elements and others, containing a bool-ish
    string in its ``val`` attribute, xsd:boolean plus 'on' and 'off'.
    """
    # val = OptionalAttribute('w:val', ST_String, default=True)
    tag = ZeroOrOne('w:tag')
    alias = ZeroOrOne('w:alias')

class CT_Block(BaseOxmlElement):
    """
    Used for ``<w:b>``, ``<w:i>`` elements and others, containing a bool-ish
    string in its ``val`` attribute, xsd:boolean plus 'on' and 'off'.
    """
    # val = OptionalAttribute('w:val', ST_String, default=True)
    properties = ZeroOrOne('w:sdtPr', successors=())
    content = ZeroOrOne('w:sdtContent', successors=())

    def hasTag(self, tag):
        if self.properties != None:
            if self.properties.tag != None:
                if self.properties.tag.val == tag:
                    return True
        return False
    
    def getContent(self):
        return self.content


class CT_BlockContent(BaseOxmlElement):
    """
    Used for ``<w:b>``, ``<w:i>`` elements and others, containing a bool-ish
    string in its ``val`` attribute, xsd:boolean plus 'on' and 'off'.
    """
    p = ZeroOrMore('w:p')
    sdt = ZeroOrMore('w:sdt')
    r = ZeroOrMore('w:r')

    def hasTag(self, tag):
        for s in self.sdt_lst:
            if s.hasTag(tag) == True:
                return True
                # if s.getContent().hasField(tag):
                #     return True

        for p in self.p_lst:
            for s in p.sdt_lst:
                if s.hasTag(tag) == True:
                    return True
                    # if s.getContent().hasField(tag):
                    #     return True
        return False

    def getTag(self, tag):
        for s in self.sdt_lst:
            if s.hasTag(tag) == True:
                return s.getContent()

        for p in self.p_lst:
            for s in p.sdt_lst:
                if s.hasTag(tag) == True:
                    return s.getContent()
                # if s.getContent().hasField(tag):
                #     return s.getContent().getContent()
        return None

    def hasField(self, field):
        if self.hasTag('fields'):
            c = self.getTag('fields')
            if c.hasTag(field):
                return True
        return False

    def getField(self, field):
        if self.hasTag('fields'):
            c = self.getTag('fields')
            if c.hasTag(field):
                return c.getTag(field)
        return None


    # def hasFields(self):
    #     count = 0
    #     for par in self.p_lst:
    #         count += len(par.getFields())
    #     return count > 0

    # def findFieldValue(self):
    #     pass

    # def getFields(self):
    #     for par in self.p_lst:
    #         for r in par.getFields():
    #             yield r


    # def hasSdtFieldTagInP(self):
    #     for par in self.p_lst:
    #         if par.hasFieldInSdt() == True:
    #             return True
    #     return False

    @property
    def all_text(self):
        res = ''
        for r in self.r_lst:
            for t in r.t_lst:
                res += t.text
        return res

class CT_P_Custom(CT_P):
    """
    ``<w:p>`` element, containing the properties and text for a paragraph.
    """
    sdt = ZeroOrMore('w:sdt')

    def getFields(self):
        for s in self.sdt_lst:
            if s.properties != None:
                if s.properties.tag != None:
                    if s.properties.tag.val == 'fields':
                        yield s.content

    def hasFieldInSdt(self):
        for s in self.sdt_lst:
            if s.properties != None:
                if s.properties.tag != None:
                    if s.properties.tag.val == 'fields':
                        print('yay!')
                        return True

        return False


# register_element_cls()
register_element_cls('w:alias', CT_Alias)
register_element_cls('w:tag', CT_Tag)
register_element_cls('w:sdt', CT_Block)
register_element_cls('w:sdtPr', CT_BlockProperties)
register_element_cls('w:sdtContent', CT_BlockContent)
register_element_cls('w:p', CT_P_Custom)


# def test3():
#     document = Document('Software Validation and Verification Plan (1).docx')
#     body = document._body._body
#     ps = body.xpath('//w:sdt')
#     for p in ps:
#         pass
#         # yield Run(p, document._body)


# def test2():
#     document = Document('test_doc.docx')
#     body = document._body._body
#     ps = body.xpath('//w:t')
#     for p in ps:
#         yield Run(p, document._body)


# def test():
#     document = Document('test_doc.docx')
#     body = document._body._body
#     ps = body.xpath('//w:p')
#     for p in ps:
#         yield Paragraph(p, document._body)

# gen = test3()
# gen = test2()
# for a in gen:
#     print(a.text)
# document = Document('test_doc.docx')
# body = document._body._body

# o = open("test_doc.docx", "rb")


# xml = '''
#  <root>
#    <element key='value'>text</element>
#    <element>text</element>tail
#    <empty-element xmlns="http://testns/" />
#  </root>
#  '''
# root = ET.fromstring(o)
# print(root)

# context = etree.iterparse(StringIO(xml))

# # context = etree.iterparse(o)
# for action, elem in context:
#     print("%s: %s" % (action, elem.tag))

document = Document('Software Validation and Verification Plan (1).docx')

doc_elm = document._element
contentControls = doc_elm.xpath('.//w:sdt/w:sdtContent')
# for contentControl in contentControls:
#     print(contentControl)

content_list = doc_elm.xpath('.//w:sdt')
for content in content_list:
    # if contentControl.properties.tag != None:

    #     print(contentControl.properties.tag.val)
    #     if contentControl.properties.tag.val == 'workItem':
    #         print('Found a workitem')
    #         if contentControl.content.hasSdtFieldTagInP():
    #             print('And a field..')
    # fields = contentControl.content.getFields()
    # for field in fields:
    #     print(field)
    if content.hasTag('workItem'):
        workitem = content.getContent()
        if workitem.hasField('id'):
            id = workitem.getField('id').all_text
            # print(id)
            par = Paragraph(workitem.p_lst[0], document._body)
            # va= 10
            par.add_run(f'\nworkitem: {id}')

document.save('out.docx')
            

    # for p in contentControl.p_lst:
    #     par = Paragraph(p, document._body)
    #     # new_run = Run()
    #     # par.runs.insert(0, )
    #     par.add_run('Sample t')
    #     print(par.text)
    #     if 'a line' in par.text:
    #         par.text += ' added content'
    # contentControl.test_pars()
    # print(f'a: {contentControl}')

# document.save('out.docx')
#     for p in contentControl.p_lst:
#         par = Paragraph(p, document._body)
#         # new_run = Run()
#         # par.runs.insert(0, )
#         par.add_run('Sample t')
#         print(par.text)
#         if 'a line' in par.text:
#             par.text += ' added content'
#     # contentControl.test_pars()
#     print(f'a: {contentControl}')

# document.save('out.docx')
    # for a in contentControl:
    #     print(f'b: {a}')
    #     for b in a:
    #         print(f'c: {b}')
    #         for c in b:
    #             print(f'd: {c}')
    #             for d in c:
    #                 print(f'e: {d}')

# for par in document.paragraphs:
#     p = par._element
#     sdts = p.xpath('w:sdt')
#     for sdt in sdts:
#         print('found an sdt element')
#     print(par.text)
#     for run in par.runs:
# print(run.text)
