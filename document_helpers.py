from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
import progressbar
import logging
import copy
import re

def findWorkitemInDoc(doc):
    workitems_in_doc = {}
    doc_elm = doc._element

    content_list = doc_elm.xpath('.//w:sdt')
    print('Finding workitems in document')
    for content in progressbar.progressbar(content_list, redirect_stdout=True):
        if content.hasTag('workItem'):
            workitem = content.getContent()
            if workitem.hasField('id'):
                id = workitem.getField('id').all_text
                workitems_in_doc[id] = None
    
    logging.info(f"Found {len(workitems_in_doc)} in source document")
    return workitems_in_doc



def matchResultsToDoc(workitems, test_runs):
    for test_run in test_runs:
        print(f'Mathing test run "{test_run}" results to document')
        for record in progressbar.progressbar(test_runs[test_run].records, redirect_stdout=True):
            if record.getTestCaseName() in workitems:
                workitems[record.getTestCaseName()] = record
            else:
                logging.warning(f'{record.getTestCaseName()} was in test run {test_run} but not in the document')
                print(f'{record.getTestCaseName()} was in test run {test_run} but not in the document')
    for workitem in workitems:
        if workitems[workitem] == None:
            logging.warning(f'{workitem} not found in test runs')
            print(f'{workitem} not found in test runs')

def extendPolarionTables(doc):
    doc_elm = doc._element

    content_list = doc_elm.xpath('.//w:sdt')
    print('preparing document')
    for content in progressbar.progressbar(content_list, redirect_stdout=True):
        if content.hasTag('workItem'):
            workitem = content.getContent()
            id = ''
            if workitem.hasField('id'):
                id = workitem.getField('id').all_text

            if workitem.hasTag('_internal_testSteps'):
                workitem_content = workitem.getTag('_internal_testSteps')
                
                try:
                    for tr in workitem_content.tbl_lst[0].tr_lst:
                        new_tc = copy.deepcopy(tr.tc_lst[-1])
                        for p in new_tc.p_lst:
                            for r in p.r_lst:
                                r.text = ''
                        tr.tc_lst[-1].addnext(new_tc)

                    workitem_content.tbl_lst[0].tr_lst[0].tc_lst[-1].p_lst[0].r_lst[0].text = 'Result'
                except:
                    logging.error(f'Could not extend table ({id})')
                    print(f'Could not extend table ({id})')

def fillPolarionTables(doc, workitem_list, config):
    doc_elm = doc._element

    content_list = doc_elm.xpath('.//w:sdt')
    print('filling document tables')
    for content in progressbar.progressbar(content_list, redirect_stdout=True):
        if content.hasTag('workItem'):
            workitem = content.getContent()
            id = ''
            if workitem.hasField('id'):
                id = workitem.getField('id').all_text

            if workitem.hasTag('_internal_testSteps'):
                workitem_content = workitem.getTag('_internal_testSteps')

                for i, tr in enumerate(workitem_content.tbl_lst[0].tr_lst):
                    #skip header
                    if i != 0 and len(tr.tc_lst[-1].p_lst) > 0:
                        par = Paragraph(tr.tc_lst[-1].p_lst[-1], doc._body)
                        makeTestStepResult(i, par, id, workitem_list, config)
                        
def fillDocWithResults(doc, workitems, config):
    doc_elm = doc._element

    content_list = doc_elm.xpath('.//w:sdt')
    print('Filling out result in document')
    for content in progressbar.progressbar(content_list, redirect_stdout=True):
        if content.hasTag('workItem'):
            workitem = content.getContent()
            if workitem.hasField('id'):
                id = workitem.getField('id').all_text

                result = workitems[id]
                if result != None:
                    par_idx = 0
                    if config['result_position'] < len(workitem.p_lst):
                        par_idx = config['result_position']
                    else:
                        logging.error(f'Position failure for {id}')
                        print(f'Position failure for {id}')

                    par = Paragraph(workitem.p_lst[par_idx], doc._body)
                    makeTestCaseResult(par, id, workitems, config)
                    
def cleanhtml(raw_html):
    cleanr = re.compile('<.*?>')
    cleantext = re.sub(cleanr, '', raw_html)
    return cleantext

def getResultColor(result, config):
    if 'result_name_color' in config:
        if result in config['result_name_color']:
            config_color = config['result_name_color'][result]
            return RGBColor(config_color[0], config_color[1], config_color[2])
    return RGBColor(0, 0, 0)

def makeTestStepResult(step_id, paragraph, workitem_id, workitem_list, config):
    w = workitem_list[workitem_id]

    result = None
    comment = None
    attachments = None

    if w.testStepResults != None:
        if len(w.testStepResults.TestStepResult) >= step_id:
            if w.testStepResults.TestStepResult[step_id - 1].result != None:
                result = w.testStepResults.TestStepResult[step_id - 1].result.id
            if w.testStepResults.TestStepResult[step_id - 1].comment != None:
                comment = w.testStepResults.TestStepResult[step_id - 1].comment.content
            if w.testStepResults.TestStepResult[step_id - 1].attachments != None:
                attachments = True

    if result != None:
        run = paragraph.add_run(f'{result.upper()}\n')        
        run.font.color.rgb = getResultColor(result, config)
        run.font.bold = True
    if comment != None:
        paragraph.add_run(f'{cleanhtml(comment)}\n')
    if attachments != None:
        paragraph.add_run(f'Has attachments')


def makeTestCaseResult(paragraph, workitem_id, workitem_list, config):
    result = workitem_list[workitem_id]
    result_string = config['result_string']
    if result.executed == None:
        result_string = result_string.replace('{result}', '-')
        result_string = result_string.replace('{result_color}', '-')
        result_string = result_string.replace('{executed}', '-')
        result_string = result_string.replace('{user}', '-')
        result_string = result_string.replace('{comment}', '-')
    else:
        result_string = result_string.replace('{result}', result.result.id)
        # result.executedByURI
        result_string = result_string.replace('{executed}', result.executed.strftime(config['date_format'])) 
        result_string = result_string.replace('{user}', result.getExecutingUser().name)
        if result.comment == None:
            result_string = result_string.replace('{comment}', '-')
        else:
            # result.comment.content can contain HTML, make sure to clean it before putting in doc
            comment = cleanhtml(result.comment.content)
            result_string = result_string.replace('{comment}', comment)

    if '{result_color}' in result_string:
        result_color = result.result.id.upper()

        parts = result_string.split('{result_color}')
        paragraph.add_run(parts[0])
        result_run = paragraph.add_run(result_color)
        paragraph.add_run(parts[1])

        result_run.font.color.rgb = getResultColor(result.result.id, config)
        result_run.font.bold = True
        pass
    else:

        paragraph.add_run(f'\n{result_string}')
